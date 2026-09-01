# -*- coding: utf-8 -*-
"""
Conversor de los informes HTML del proyecto a Word.

Por qué existe
--------------
Los capítulos y anexos se generan como HTML con el estilo de la casa
(`informe_estilo.py`) y no pasan por Markdown, así que `md_a_docx.py` no podía
convertirlos: solo el informe por comunidad tenía versión Word, porque él sí
escribe un `.md` intermedio. El sociólogo pidió el paquete completo en Word
(JAVIKO, 1-sep-2026), y hacerlo capítulo por capítulo habría significado
tocar seis generadores. Este conversor lo resuelve en un solo lugar y sirve
también para los informes que se escriban en el futuro.

No sustituye al HTML: los dos formatos se entregan. El HTML se abre en
cualquier equipo y se imprime mejor; el Word es el que el sociólogo edita.

Qué entiende
------------
El estilo de la casa es un conjunto cerrado y pequeño: `header`, `h1`–`h3`,
párrafos, tablas, listas, los recuadros `corte`/`nota`/`alerta`/`hallazgo`,
las tarjetas `kpis` y el `footer`. Dentro del texto respeta negrita, cursiva
y código. Las barras de porcentaje (`div.barra`) NO se dibujan: se convierte
su cifra a texto, que es la información que llevan.

Las imágenes embebidas en base64 (los mapas del informe por comunidad) se
incrustan en el Word a tamaño de página.

Uso
---
    python scripts/html_a_docx.py docs/CAPITULO-perfil-del-titular.html
    python scripts/html_a_docx.py docs/*.html --salida-dir build_entrega/word
    python scripts/html_a_docx.py --paquete      # los del paquete de entrega

Se corre con el Python del PATH (el que tiene python-docx), no con el de
OSGeo4W.
"""

import argparse
import base64
import glob
import io
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from md_a_docx import AZUL, GRIS, sombrear, texto_con_formato  # noqa: E402

BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')

# Recuadros del estilo de la casa → color de fondo en Word. Los mismos tonos
# que el CSS de informe_estilo.py, para que las dos versiones se reconozcan.
RECUADROS = {
    'corte': ('FFF8E6', '8A6100'),
    'nota': ('F0F6FF', '1E4D8C'),
    'hallazgo': ('F4FBF6', '1D6B35'),
    'alerta': ('FFF5F5', 'A03030'),
}

# Documentos del paquete de entrega que necesitan Word. El informe por
# comunidad no está: escribe su propio .md y su Word sale de md_a_docx.py con
# los mapas embebidos.
PAQUETE = [
    'INFORME-CONSOLIDADO-padron-regantes.html',
    'CAPITULO-conocimiento-y-gobernanza.html',
    'CAPITULO-predio-y-agua.html',
    'CAPITULO-produccion-agropecuaria.html',
    'CAPITULO-perfil-del-titular.html',
    'CAPITULO-estructura-del-padron.html',
    'CAPITULO-servicios-basicos.html',
    'ANEXO-operadores-por-comunidad.html',
]


# ─── Árbol mínimo del HTML ───────────────────────────────────────────────────

class Nodo:
    __slots__ = ('tag', 'clases', 'attrs', 'hijos', 'texto')

    def __init__(self, tag='', attrs=None, texto=None):
        self.tag = tag
        self.attrs = dict(attrs or [])
        self.clases = set(self.attrs.get('class', '').split())
        self.hijos = []
        self.texto = texto


VACIOS = {'br', 'img', 'hr', 'meta', 'link', 'input'}


class Arbol(HTMLParser):
    """Construye el árbol del documento. Ignora <style>, <script> y <head>."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.raiz = Nodo('body')
        self.pila = [self.raiz]
        self.ignorar = 0

    def handle_starttag(self, tag, attrs):
        if tag in ('style', 'script'):
            self.ignorar += 1
            return
        if self.ignorar or tag in ('html', 'head', 'body', 'meta', 'link'):
            return
        n = Nodo(tag, attrs)
        self.pila[-1].hijos.append(n)
        if tag not in VACIOS:
            self.pila.append(n)

    def handle_endtag(self, tag):
        if tag in ('style', 'script'):
            self.ignorar = max(0, self.ignorar - 1)
            return
        if self.ignorar or tag in VACIOS:
            return
        for i in range(len(self.pila) - 1, 0, -1):
            if self.pila[i].tag == tag:
                del self.pila[i:]
                break

    def handle_data(self, dato):
        if not self.ignorar and dato.strip():
            self.pila[-1].hijos.append(Nodo('#texto', texto=dato))


# ─── Texto en línea: se traduce a los marcadores que ya entiende md_a_docx ───

def inline(nodo):
    """Texto del nodo con **negrita**, *cursiva* y `código`, como Markdown.

    Así se reutiliza `texto_con_formato()` sin escribir un segundo formateador.
    """
    if nodo.texto is not None:
        return re.sub(r'\s+', ' ', nodo.texto)
    partes = []
    for h in nodo.hijos:
        if h.tag == 'br':
            partes.append(' ')
            continue
        t = inline(h)
        if not t.strip():
            partes.append(t)
            continue
        if h.tag in ('b', 'strong'):
            partes.append(f'**{t.strip()}**')
        elif h.tag in ('i', 'em'):
            # la barra de porcentaje guarda su cifra en un <i>: no es cursiva
            partes.append(t.strip() if 'barra' in nodo.clases else f'*{t.strip()}*')
        elif h.tag == 'code':
            partes.append(f'`{t.strip()}`')
        else:
            partes.append(t)
    return ''.join(partes)


def texto_de(nodo):
    return re.sub(r'\s+', ' ', inline(nodo)).strip()


# ─── Piezas de Word ──────────────────────────────────────────────────────────

def parrafo(doc, texto, tam=10.5, color=None, cursiva=False, centrado=False,
            espacio_antes=None):
    p = doc.add_paragraph()
    texto_con_formato(p, texto)
    for r in p.runs:
        r.font.size = Pt(tam)
        if color is not None:
            r.font.color.rgb = color
        if cursiva:
            r.italic = True
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if espacio_antes is not None:
        p.paragraph_format.space_before = Pt(espacio_antes)
    return p


def titulo(doc, texto, nivel):
    if nivel == 1:
        p = doc.add_paragraph()
        texto_con_formato(p, texto)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(19)
            r.font.color.rgb = AZUL
        p.paragraph_format.space_after = Pt(10)
        return
    h = doc.add_heading(level=min(nivel, 4))
    h.paragraph_format.space_before = Pt(12 if nivel == 2 else 9)
    h.paragraph_format.space_after = Pt(4)
    for r in list(h.runs):
        r.text = ''
    texto_con_formato(h, texto)
    for r in h.runs:
        r.font.color.rgb = AZUL
        r.font.size = Pt(14 if nivel == 2 else 12)
        r.bold = True


def recuadro(doc, texto, fondo, color_texto):
    """Los avisos del estilo de la casa: una celda sombreada del ancho de la
    página. Word no tiene «div», y una tabla de 1x1 es lo que mejor lo imita."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    celda = t.rows[0].cells[0]
    celda.text = ''
    texto_con_formato(celda.paragraphs[0], texto)
    for r in celda.paragraphs[0].runs:
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(color_texto)
    sombrear(celda, fondo)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tarjetas_kpi(doc, kpis):
    """div.kpis → una fila de columnas con la cifra grande y su etiqueta."""
    if not kpis:
        return
    t = doc.add_table(rows=2, cols=len(kpis))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, (numero, etiqueta) in enumerate(kpis):
        c = t.rows[0].cells[k]
        c.text = ''
        r = c.paragraphs[0].add_run(numero)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = AZUL
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sombrear(c, 'F8FAFC')
        c2 = t.rows[1].cells[k]
        c2.text = ''
        r2 = c2.paragraphs[0].add_run(etiqueta)
        r2.font.size = Pt(8)
        r2.font.color.rgb = GRIS
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sombrear(c2, 'F8FAFC')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def escribir_tabla(doc, nodo):
    filas = [f for f in recorrer(nodo, 'tr')]
    if not filas:
        return
    encabezados, cuerpo = [], []
    for fila in filas:
        celdas = [c for c in fila.hijos if c.tag in ('th', 'td')]
        valores = [texto_de(c) for c in celdas]
        if celdas and all(c.tag == 'th' for c in celdas) and not encabezados:
            encabezados = valores
        else:
            cuerpo.append((valores, 'dest' in fila.clases))
    if not encabezados:
        encabezados, cuerpo = cuerpo[0][0], cuerpo[1:]
    ancho = len(encabezados)

    t = doc.add_table(rows=1, cols=ancho)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, texto in enumerate(encabezados):
        c = t.rows[0].cells[k]
        c.text = ''
        texto_con_formato(c.paragraphs[0], texto, base_negrita=True)
        for r in c.paragraphs[0].runs:
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(c, '1F4E79')
    for valores, destacada in cuerpo:
        cs = t.add_row().cells
        for k, texto in enumerate(valores[:ancho]):
            cs[k].text = ''
            texto_con_formato(cs[k].paragraphs[0], texto,
                              base_negrita=destacada)
            for r in cs[k].paragraphs[0].runs:
                r.font.size = Pt(8.5)
            if destacada:
                sombrear(cs[k], 'EEFBF0')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def recorrer(nodo, tag):
    for h in nodo.hijos:
        if h.tag == tag:
            yield h
        else:
            yield from recorrer(h, tag)


def imagen(doc, nodo):
    src = nodo.attrs.get('src', '')
    m = re.match(r'data:image/(png|jpe?g);base64,(.+)$', src, re.S)
    if m:
        doc.add_picture(io.BytesIO(base64.b64decode(m.group(2))), width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── Recorrido del documento ─────────────────────────────────────────────────

def escribir(doc, nodo):
    for h in nodo.hijos:
        tag, clases = h.tag, h.clases

        if tag == 'header':
            for hh in recorrer(h, 'h1'):
                titulo(doc, texto_de(hh), 1)
            for hh in h.hijos:
                for sub in recorrer(hh, 'p'):
                    if 'sub' in sub.clases:
                        parrafo(doc, texto_de(sub), 9.5, GRIS)
            for hh in recorrer(h, 'div'):
                if 'meta' in hh.clases:
                    parrafo(doc, texto_de(hh), 8, GRIS)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

        elif tag in ('h1', 'h2', 'h3', 'h4'):
            titulo(doc, texto_de(h), int(tag[1]))

        elif tag == 'p':
            t = texto_de(h)
            if t:
                if 'sub' in clases:
                    parrafo(doc, t, 8.5, GRIS, cursiva=True)
                else:
                    parrafo(doc, t)

        elif tag == 'table':
            escribir_tabla(doc, h)

        elif tag in ('ul', 'ol'):
            for li in h.hijos:
                if li.tag != 'li':
                    continue
                p = doc.add_paragraph(style='List Bullet' if tag == 'ul'
                                      else 'List Number')
                texto_con_formato(p, texto_de(li))
                for r in p.runs:
                    r.font.size = Pt(10.5)

        elif tag == 'img':
            imagen(doc, h)

        elif tag == 'footer':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            from md_a_docx import linea_horizontal
            linea_horizontal(p)
            # el pie son dos <span> que en el HTML se separan por el layout;
            # sin unirlos a mano salían pegados («PorotogCorte: 5 de agosto»)
            trozos = [texto_de(s) for s in h.hijos if s.tag == 'span']
            parrafo(doc, ' · '.join(t for t in trozos if t) or texto_de(h),
                    8, GRIS, centrado=True)

        elif tag == 'div':
            aviso = next((c for c in clases if c in RECUADROS), None)
            if aviso:
                fondo, color = RECUADROS[aviso]
                recuadro(doc, texto_de(h), fondo, color)
            elif 'kpis' in clases:
                kpis = []
                for k in h.hijos:
                    if 'kpi' not in k.clases:
                        continue
                    n = next((texto_de(x) for x in k.hijos if 'n' in x.clases), '')
                    t = next((texto_de(x) for x in k.hijos if 't' in x.clases), '')
                    kpis.append((n, t))
                tarjetas_kpi(doc, kpis)
            elif 'barra' in clases:
                continue  # su cifra ya va en el texto de la celda
            else:
                escribir(doc, h)

        elif tag == '#texto':
            pass
        else:
            escribir(doc, h)


def convertir(ruta_html, ruta_docx):
    with open(ruta_html, encoding='utf-8') as f:
        arbol = Arbol()
        arbol.feed(f.read())

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    # las tablas de estos informes son anchas: sin esto Word las desborda
    ajustes = doc.settings.element
    if ajustes.find(qn('w:doNotAutoCompressPictures')) is None:
        ajustes.append(OxmlElement('w:doNotAutoCompressPictures'))

    escribir(doc, arbol.raiz)
    doc.save(ruta_docx)
    return len(doc.tables), len(doc.paragraphs)


def main():
    ap = argparse.ArgumentParser(description=__doc__.split('\n')[1])
    ap.add_argument('archivos', nargs='*', help='HTML a convertir')
    ap.add_argument('--paquete', action='store_true',
                    help='convierte los documentos del paquete de entrega')
    ap.add_argument('--salida-dir', default=None,
                    help='carpeta de salida (por defecto, junto al HTML)')
    args = ap.parse_args()

    rutas = []
    if args.paquete:
        rutas = [os.path.join(BASE, 'docs', n) for n in PAQUETE]
    for a in args.archivos:
        rutas.extend(sorted(glob.glob(a)) or [a])
    if not rutas:
        ap.error('no se indicó ningún archivo (o usa --paquete)')

    for ruta in rutas:
        if not os.path.exists(ruta):
            print(f'  ⚠ no existe: {ruta}')
            continue
        destino_dir = args.salida_dir or os.path.dirname(ruta)
        os.makedirs(destino_dir, exist_ok=True)
        destino = os.path.join(
            destino_dir,
            os.path.splitext(os.path.basename(ruta))[0] + '.docx')
        tablas, parrafos = convertir(ruta, destino)
        kb = os.path.getsize(destino) / 1024
        print(f'  {os.path.basename(ruta)}  ->  {os.path.basename(destino)}'
              f'  ({tablas} tablas, {parrafos} párrafos, {kb:,.0f} KB)')


if __name__ == '__main__':
    main()
