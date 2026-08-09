# -*- coding: utf-8 -*-
"""
Conversor de Markdown a Word, para los informes que salen del proyecto.

Por qué existe
--------------
Los informes se escriben y se versionan en Markdown (se leen en el repositorio,
se comparan entre versiones y no dependen de ningún programa). Pero lo que se
envía a la dirección del proyecto o al consorcio tiene que ser un `.docx`.

No hay pandoc en esta máquina, así que la conversión se hace con `python-docx`,
que sí está en el Python del sistema (es el que ya usa el proyecto para sus
otros informes en Word).

Qué entiende del Markdown
-------------------------
Lo que usan los informes de este proyecto: títulos, párrafos, listas, tablas,
citas, líneas separadoras, y dentro del texto **negrita**, *cursiva* y `código`.
No pretende ser un conversor general.

Tablas muy largas
-----------------
`--max-filas N` corta las tablas que pasen de N filas y deja una nota en su
lugar. Hace falta para el documento de revisión de campo: sus listados suman
más de 9.000 filas y en Word darían cientos de páginas que nadie va a leer.
Ese detalle se trabaja en el Excel; el Word es para entender y decidir.

Uso
---
    python scripts/md_a_docx.py docs/INFORME-cartografia-consorcio.md
    python scripts/md_a_docx.py docs/REVISION-CAMPO.md --max-filas 15 \
        --salida "docs/GUIA-REVISION-CAMPO.docx"
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)


def sombrear(celda, color_hex):
    tc = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc.append(shd)


def texto_con_formato(parrafo, texto, base_negrita=False):
    """Escribe el texto respetando **negrita**, *cursiva* y `código`."""
    # se corta por los marcadores conservándolos, para reconstruir el formato
    partes = re.split(r'(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)', texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
            r = parrafo.add_run(parte[2:-2])
            r.bold = True
        elif parte.startswith('`') and parte.endswith('`') and len(parte) > 2:
            r = parrafo.add_run(parte[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8B, 0x25, 0x00)
        elif (parte.startswith('*') and parte.endswith('*')
              and len(parte) > 2 and not parte.startswith('**')):
            r = parrafo.add_run(parte[1:-1])
            r.italic = True
        else:
            # el marcador de salto se convierte aquí en salto real de línea
            piezas = parte.split(SALTO)
            r = parrafo.add_run(piezas[0])
            for pieza in piezas[1:]:
                r.add_break()
                r = parrafo.add_run(pieza)
        if base_negrita:
            r.bold = True


def es_separador_tabla(linea):
    return bool(re.match(r'^\|[\s:|-]+\|$', linea.strip()))


def celdas(linea):
    return [c.strip() for c in linea.strip().strip('|').split('|')]


def linea_horizontal(parrafo):
    p = parrafo._p.get_or_add_pPr()
    borde = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'BFBFBF')
    borde.append(bottom)
    p.append(borde)


def continuar_bloque(lineas, i, primera):
    """
    Junta las líneas que continúan el mismo párrafo o ítem de lista.

    Devuelve (trozos, saltos). Cada trozo es (texto, salto_forzado): en Markdown
    dos espacios al final de una línea significan «salto de línea aquí», y se
    usan en las cabeceras de los informes (Proyecto / Asunto / Fecha). Sin
    respetarlos, esas tres líneas salían corridas en un solo renglón.
    """
    trozos = [(primera.strip(), lineas[i].endswith('  '))]
    saltos = 0
    j = i
    while (j + 1 < len(lineas) and lineas[j + 1].strip()
           and not re.match(r'^\s*([-*·]\s|\d+\.\s|#{1,4}\s|>|\|)', lineas[j + 1])
           and lineas[j + 1].strip() not in ('---', '***', '___')):
        j += 1
        saltos += 1
        trozos.append((lineas[j].strip(), lineas[j].endswith('  ')))
    return trozos, saltos


SALTO = '\x00'      # marcador interno de «aquí va un salto de línea»


def escribir_bloque(parrafo, trozos):
    """
    Une el bloque en un solo texto y ENTONCES le aplica el formato.

    Importa el orden: una negrita puede empezar en una línea del Markdown y
    cerrarse en la siguiente. Formateando línea por línea, esos `**` quedaban
    sin pareja y se imprimían tal cual en el documento final.
    """
    partes = []
    for k, (texto, salto) in enumerate(trozos):
        if k:
            partes.append(SALTO if trozos[k - 1][1] else ' ')
        partes.append(texto)
    texto_con_formato(parrafo, ''.join(partes))
    return SALTO in ''.join(partes)


def convertir(ruta_md, ruta_docx, max_filas=0):
    with open(ruta_md, encoding='utf-8') as f:
        lineas = f.read().split('\n')

    doc = Document()
    for seccion in doc.sections:
        seccion.top_margin = Cm(2.2)
        seccion.bottom_margin = Cm(2.2)
        seccion.left_margin = Cm(2.4)
        seccion.right_margin = Cm(2.4)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    i = 0
    tablas_cortadas = 0
    while i < len(lineas):
        linea = lineas[i].rstrip()

        if not linea.strip():
            i += 1
            continue

        # ── separador ──
        if linea.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            linea_horizontal(p)
            i += 1
            continue

        # ── títulos ──
        m = re.match(r'^(#{1,4})\s+(.*)$', linea)
        if m:
            nivel = len(m.group(1))
            texto = m.group(2).strip()
            if nivel == 1:
                p = doc.add_paragraph()
                r = p.add_run(texto)
                r.bold = True
                r.font.size = Pt(19)
                r.font.color.rgb = AZUL
                p.paragraph_format.space_after = Pt(12)
            else:
                h = doc.add_heading(level=min(nivel, 4))
                h.paragraph_format.space_before = Pt(12 if nivel == 2 else 9)
                h.paragraph_format.space_after = Pt(4)
                for r in list(h.runs):
                    r.text = ''
                texto_con_formato(h, texto)
                for r in h.runs:
                    r.font.color.rgb = AZUL
                    r.font.size = Pt(14 if nivel == 2 else 12)
                    r.bold = True
            i += 1
            continue

        # ── tabla ──
        if linea.strip().startswith('|') and i + 1 < len(lineas) \
                and es_separador_tabla(lineas[i + 1]):
            encabezados = celdas(linea)
            filas = []
            j = i + 2
            while j < len(lineas) and lineas[j].strip().startswith('|'):
                filas.append(celdas(lineas[j]))
                j += 1

            recortada = False
            if max_filas and len(filas) > max_filas:
                filas = filas[:max_filas]
                recortada = True
                tablas_cortadas += 1

            tabla = doc.add_table(rows=1, cols=len(encabezados))
            tabla.style = 'Table Grid'
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            for k, texto in enumerate(encabezados):
                celda = tabla.rows[0].cells[k]
                celda.text = ''
                texto_con_formato(celda.paragraphs[0], texto, base_negrita=True)
                for r in celda.paragraphs[0].runs:
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                sombrear(celda, '1F4E79')
            for fila in filas:
                cs = tabla.add_row().cells
                for k, texto in enumerate(fila[:len(encabezados)]):
                    cs[k].text = ''
                    texto_con_formato(cs[k].paragraphs[0], texto)
                    for r in cs[k].paragraphs[0].runs:
                        r.font.size = Pt(9.5)
            if recortada:
                p = doc.add_paragraph()
                r = p.add_run('Tabla recortada: el listado completo está en '
                              'REVISION-CAMPO.xlsx, donde además se puede '
                              'filtrar y ordenar.')
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = GRIS
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i = j
            continue

        # ── cita ──
        if linea.strip().startswith('>'):
            texto = re.sub(r'^\s*>\s?', '', linea)
            while i + 1 < len(lineas) and lineas[i + 1].strip().startswith('>'):
                i += 1
                texto += ' ' + re.sub(r'^\s*>\s?', '', lineas[i]).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(4)
            texto_con_formato(p, texto.strip())
            for r in p.runs:
                r.font.color.rgb = GRIS
                if not r.bold:
                    r.italic = True
            i += 1
            continue

        # ── listas ──
        # Un ítem puede seguir en las líneas de abajo (Markdown lo permite); si
        # no se recogen aquí, el resto del ítem sale como un párrafo suelto
        # debajo de la viñeta, que es como quedaba antes.
        m = re.match(r'^\s*([-*·]|\d+\.)\s+(.*)$', linea)
        if m:
            estilo = 'List Bullet' if m.group(1) in ('-', '*', '·') else 'List Number'
            bloque, saltos = continuar_bloque(lineas, i, m.group(2))
            p = doc.add_paragraph(style=estilo)
            p.paragraph_format.space_after = Pt(3)
            escribir_bloque(p, bloque)
            i += 1 + saltos
            continue

        # ── párrafo ──
        bloque, saltos = continuar_bloque(lineas, i, linea.strip())
        p = doc.add_paragraph()
        # Justificar un párrafo con saltos forzados estira la línea corta hasta
        # el margen y deja huecos enormes entre palabras (pasaba en la cabecera
        # Proyecto / Asunto / Fecha).
        if not escribir_bloque(p, bloque):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1 + saltos

    doc.save(ruta_docx)
    return tablas_cortadas


def main():
    ap = argparse.ArgumentParser(description='Convierte un informe .md a .docx')
    ap.add_argument('entrada')
    ap.add_argument('--salida', default=None)
    ap.add_argument('--max-filas', type=int, default=0,
                    help='corta las tablas que pasen de N filas (0 = sin corte)')
    args = ap.parse_args()

    entrada = os.path.abspath(args.entrada)
    if not os.path.exists(entrada):
        print("ERROR: no existe {}".format(entrada))
        return 1
    salida = os.path.abspath(args.salida) if args.salida \
        else os.path.splitext(entrada)[0] + '.docx'

    cortadas = convertir(entrada, salida, args.max_filas)
    print("  {}  ->  {}  ({:,.0f} KB)".format(
        os.path.basename(entrada), os.path.basename(salida),
        os.path.getsize(salida) / 1024))
    if cortadas:
        print("     {} tabla(s) recortada(s) por superar {} filas"
              .format(cortadas, args.max_filas))
    return 0


if __name__ == '__main__':
    sys.exit(main())
